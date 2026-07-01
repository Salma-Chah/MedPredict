# -*- coding: utf-8 -*-
"""
Genere rapport_technique.docx depuis rapport_technique.md
Dependance : python-docx (pip install python-docx)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rapport_technique.md")
DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rapport_technique.docx")

# ── Palette couleurs (alignee sur generate_pdf.py) ────────────────────────
BLUE_DARK   = RGBColor(0x1E, 0x40, 0x57)
BLUE_MID    = RGBColor(0x1A, 0x52, 0x76)
BLUE_LIGHT  = "D6EAF8"
GREEN_DARK  = RGBColor(0x1E, 0x84, 0x49)
CODE_BG     = "F5F5F5"
CODE_BORDER = "B4B4B4"
GRAY_TEXT   = RGBColor(0x5A, 0x5A, 0x5A)
BLACK       = RGBColor(0x1E, 0x1E, 0x1E)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
SEP_COLOR   = "B4C8DC"
NOTE_BG     = "FCF3CF"
NOTE_BORDER = "F1C40F"

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


# ── Helpers XML (ombrage, bordures) ───────────────────────────────────────

def shade_pPr(pPr, color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def left_border_pPr(pPr, color_hex, size=24):
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def box_border_paragraph(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color_hex)
        pBdr.append(el)
    pPr.append(pBdr)


def bottom_border_paragraph(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_paragraph(paragraph, color_hex):
    shade_pPr(paragraph._p.get_or_add_pPr(), color_hex)


def left_border_paragraph(paragraph, color_hex, size=24):
    left_border_pPr(paragraph._p.get_or_add_pPr(), color_hex, size)


def shade_style(style, color_hex):
    shade_pPr(style.element.get_or_add_pPr(), color_hex)


def left_border_style(style, color_hex, size=24):
    left_border_pPr(style.element.get_or_add_pPr(), color_hex, size)


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


# ── Texte / markdown inline ────────────────────────────────────────────────

def strip_md(text: str) -> str:
    """Retire les balises markdown inline (** * `)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    return text


def add_inline_runs(paragraph, text, italic=False):
    """Decoupe **gras**, *italique* et `code` et ajoute les runs correspondants."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token)
        if italic:
            run.italic = True


# ── Blocs de contenu ────────────────────────────────────────────────────────

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    bottom_border_paragraph(p, SEP_COLOR)


def add_quote(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    shade_paragraph(p, NOTE_BG)
    left_border_paragraph(p, NOTE_BORDER, size=24)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(lines):
        if idx > 0:
            p.add_run().add_break()
        add_inline_runs(p, line, italic=True)


def add_code_block(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    shade_paragraph(p, CODE_BG)
    box_border_paragraph(p, CODE_BORDER)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.2)
    for idx, line in enumerate(lines):
        run = p.add_run(line if line.strip() else " ")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x32, 0x32, 0x32)
        if idx < len(lines) - 1:
            p.add_run().add_break()


def add_table(doc, rows):
    if len(rows) < 2:
        return

    def parse_row(line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    header_cells = parse_row(rows[0])
    data_rows = [parse_row(r) for r in rows[2:] if r.strip()]
    n_cols = len(header_cells)
    if n_cols == 0:
        return

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for j, text in enumerate(header_cells):
        cell = hdr_cells[j]
        set_cell_background(cell, "1E4057")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(strip_md(text))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)

    for i, row in enumerate(data_rows):
        cells = table.add_row().cells
        bg = "FFFFFF" if i % 2 == 0 else BLUE_LIGHT
        for j in range(n_cols):
            text = row[j] if j < len(row) else ""
            cell = cells[j]
            set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            add_inline_runs(para, text)
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()


# ── Page de garde ─────────────────────────────────────────────────────────

def build_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    shade_paragraph(p, "1E4057")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = WHITE

    p2 = doc.add_paragraph()
    shade_paragraph(p2, "1E4057")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(50)
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0xBE, 0xDC, 0xF0)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(desc, BLUE_LIGHT)
    box_border_paragraph(desc, "AED6F1")
    desc.paragraph_format.space_before = Pt(6)
    desc.paragraph_format.space_after = Pt(6)
    run = desc.add_run(
        "Ce document presente l'ensemble du travail realise dans le cadre "
        "du projet de fin d'annee : nettoyage des donnees, entrainement "
        "des modeles de Machine Learning, deploiement via FastAPI et "
        "Streamlit, ainsi que la conception UML du systeme."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    infos = [
        ("Dataset", "hospital_readmissions_30k.csv - 30 000 patients"),
        ("Modeles", "Logistic Regression / Random Forest / XGBoost"),
        ("Outils",  "Python, scikit-learn, MLflow, FastAPI, Streamlit, SHAP"),
        ("Auteur",  "Salma Chah"),
        ("Date",    "Avril 2026"),
    ]
    for label, value in infos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label} : ")
        r1.bold = True
        r1.font.color.rgb = BLUE_DARK
        r1.font.size = Pt(10.5)
        r2 = p.add_run(value)
        r2.font.size = Pt(10.5)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    shade_paragraph(foot, "1E4057")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(40)
    run = foot.add_run("PFA 2025-2026  |  Classification binaire supervisee")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE

    doc.add_page_break()


# ── Mise en page des styles ───────────────────────────────────────────────

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = BLUE_MID
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.left_indent = Pt(8)
    left_border_style(h1, "1A5276", size=24)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = GREEN_DARK
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    bullet = doc.styles['List Bullet']
    bullet.font.name = 'Calibri'
    bullet.font.size = Pt(10.5)


# ── Parseur Markdown -> DOCX ─────────────────────────────────────────────────

def render_body(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    table_buf = []
    quote_buf = []
    code_buf = []
    in_code = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # Blocs de code ```
        if line.startswith("```"):
            if not in_code:
                if table_buf:
                    add_table(doc, table_buf)
                    table_buf = []
                if quote_buf:
                    add_quote(doc, quote_buf)
                    quote_buf = []
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, code_buf)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # Tableaux
        if line.startswith("|"):
            if quote_buf:
                add_quote(doc, quote_buf)
                quote_buf = []
            table_buf.append(line)
            i += 1
            continue
        elif table_buf:
            add_table(doc, table_buf)
            table_buf = []

        # Citations >
        if line.startswith(">"):
            quote_buf.append(line[1:].strip())
            i += 1
            continue
        elif quote_buf:
            add_quote(doc, quote_buf)
            quote_buf = []

        # Titres
        if line.startswith("#### "):
            doc.add_heading(strip_md(line[5:]), level=2)
        elif line.startswith("### "):
            doc.add_heading(strip_md(line[4:]), level=2)
        elif line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), level=1)
        elif re.match(r'^# ', line):
            pass  # deja sur la page de garde

        # Separateurs ---
        elif re.match(r'^-{3,}$', line):
            add_separator(doc)

        # Listes numerotees
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_inline_runs(p, line)

        # Listes a puces
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, line[2:])

        # Lignes vides
        elif line == "":
            pass

        # Paragraphes normaux
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)

        i += 1

    if table_buf:
        add_table(doc, table_buf)
    if quote_buf:
        add_quote(doc, quote_buf)
    if code_buf:
        add_code_block(doc, code_buf)


# ── Main ──────────────────────────────────────────────────────────────────

def main():
    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    lines = md_text.splitlines()
    title = lines[0][2:].strip()     # "# Titre"
    subtitle = lines[1][3:].strip()  # "## Sous-titre"

    start = 0
    for idx, l in enumerate(lines):
        if l.strip().startswith("---") and idx > 2:
            start = idx + 1
            break
    body_text = "\n".join(lines[start:])

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    setup_styles(doc)
    build_cover(doc, title, subtitle)
    render_body(doc, body_text)

    doc.save(DOCX_PATH)
    print(f"[OK] DOCX genere -> {DOCX_PATH}")


if __name__ == "__main__":
    main()
